import os
import io
import html
import tempfile
import streamlit as st


# conversion functions

def convert_image(input_path, output_format):
    from PIL import Image

    img = Image.open(input_path)

    if output_format.upper() in ("JPEG", "JPG") and img.mode in ("RGBA", "P", "LA"):
        background = Image.new("RGB", img.size, (255, 255, 255))
        alpha = None
        try:
            if "A" in img.getbands():
                alpha = img.getchannel("A")
        except Exception:
            alpha = None

        if alpha is not None:
            background.paste(img, mask=alpha)
        else:
            background.paste(img)

        img = background

    base = os.path.splitext(input_path)[0]
    ext = "jpg" if output_format.upper() == "JPEG" else output_format.lower()
    output_path = f"{base}_converted.{ext}"

    img.save(output_path, format=output_format.upper())
    return output_path


def convert_document(input_path, output_format):
    import_ext = os.path.splitext(input_path)[1].lower()
    base = os.path.splitext(input_path)[0]
    output_path = f"{base}_converted.{output_format.lower()}"

    text = ""

    if import_ext == ".docx":
        from docx import Document
        doc = Document(input_path)
        text = "\n".join(para.text for para in doc.paragraphs)

    elif import_ext == ".txt":
        with open(input_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

    elif import_ext == ".pdf":
        import PyPDF2
        with open(input_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            extracted_pages = []
            for page in reader.pages:
                extracted_pages.append(page.extract_text() or "")
            text = "\n".join(extracted_pages)

    elif import_ext == ".epub":
        import ebooklib
        from ebooklib import epub
        from bs4 import BeautifulSoup

        book = epub.read_epub(input_path)
        chapters = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                chapter_text = soup.get_text("\n", strip=True)
                if chapter_text:
                    chapters.append(chapter_text)

        text = "\n\n".join(chapters).strip()

    else:
        raise ValueError(f"Unsupported source format: {import_ext}")

    if output_format.lower() == "txt":
        with open(output_path, "w", encoding="utf-8", errors="replace") as f:
            f.write(text)

    elif output_format.lower() == "docx":
        from docx import Document
        new_doc = Document()
        lines = text.splitlines() or [text]
        for line in lines:
            new_doc.add_paragraph(line)
        new_doc.save(output_path)

    elif output_format.lower() == "pdf":
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Helvetica", size=11)

        lines = text.splitlines() or [text]
        for line in lines:
            safe_line = line.encode("latin-1", errors="replace").decode("latin-1")
            pdf.multi_cell(0, 7, safe_line if safe_line else " ")

        pdf.output(output_path)

    elif output_format.lower() == "epub":
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier(f"converted-{os.path.basename(base)}")
        book.set_title(os.path.basename(base) or "Converted Book")
        book.set_language("en")

        escaped_text = html.escape(text).replace("\n", "<br/>")
        chapter = epub.EpubHtml(title="Content", file_name="content.xhtml", lang="en")
        chapter.content = f"""
        <html>
          <head><title>Converted</title></head>
          <body>
            <h1>Converted</h1>
            <p>{escaped_text}</p>
          </body>
        </html>
        """

        book.add_item(chapter)
        book.toc = (chapter,)
        book.spine = ["nav", chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        epub.write_epub(output_path, book)

    else:
        raise ValueError(f"Unsupported output format: {output_format}")

    return output_path


def convert_audio_video(input_path, output_format):
    import moviepy.editor as mp

    base = os.path.splitext(input_path)[0]
    output_path = f"{base}_converted.{output_format.lower()}"
    fmt = output_format.lower()

    audio_formats = {"mp3", "wav", "ogg", "aac", "flac", "m4a"}
    video_formats = {"mp4", "avi", "mov", "mkv", "webm", "gif"}

    src_ext = os.path.splitext(input_path)[1].lower().lstrip(".")

    if src_ext in audio_formats:
        clip = mp.AudioFileClip(input_path)
        try:
            clip.write_audiofile(output_path)
        finally:
            clip.close()

    elif src_ext in video_formats:
        clip = mp.VideoFileClip(input_path)
        try:
            if fmt in audio_formats:
                if clip.audio is None:
                    raise RuntimeError("This video has no audio track to extract.")
                clip.audio.write_audiofile(output_path)
            elif fmt == "gif":
                clip.write_gif(output_path, fps=10)
            else:
                clip.write_videofile(output_path)
        finally:
            if clip.audio is not None:
                try:
                    clip.audio.close()
                except Exception:
                    pass
            clip.close()
    else:
        raise ValueError(f"Unrecognised source format: .{src_ext}")

    return output_path


IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff"}
DOC_EXTS   = {".docx", ".txt", ".pdf", ".epub"}
AUDIO_EXTS = {".mp3", ".wav", ".ogg", ".aac", ".flac", ".m4a"}
VIDEO_EXTS = {".mp4", ".avi", ".mov", ".mkv", ".webm", ".gif"}

IMAGE_OUTPUTS = ["PNG", "JPEG", "WEBP", "BMP", "GIF"]
DOC_OUTPUTS   = {
    ".docx": ["TXT", "PDF", "EPUB"],
    ".txt":  ["DOCX", "PDF", "EPUB"],
    ".pdf":  ["TXT", "DOCX", "EPUB"],
    ".epub": ["TXT", "DOCX", "PDF"],
}
AUDIO_OUTPUTS = ["MP3", "WAV", "OGG", "AAC", "FLAC"]
VIDEO_OUTPUTS = ["MP4", "AVI", "MOV", "MKV", "WEBM", "GIF", "MP3", "WAV"]


def get_output_formats(filename):
    ext = os.path.splitext(filename)[1].lower()

    if ext in IMAGE_EXTS:
        return "image", [f for f in IMAGE_OUTPUTS if f.lower() != ext.lstrip(".")]
    elif ext in DOC_EXTS:
        return "document", DOC_OUTPUTS.get(ext, [])
    elif ext in AUDIO_EXTS:
        return "audio", [f for f in AUDIO_OUTPUTS if f.lower() != ext.lstrip(".")]
    elif ext in VIDEO_EXTS:
        return "video", [f for f in VIDEO_OUTPUTS if f.lower() != ext.lstrip(".")]
    else:
        return None, []



st.set_page_config(page_title="Local File Converter", page_icon="⚡", layout="centered")

st.title("⚡ Local File Converter")
st.caption("Images · Documents · Audio/Video · EPUB support")

uploaded_file = st.file_uploader(
    "Choose a file",
    type=[
        "jpg", "jpeg", "png", "webp", "bmp", "gif", "tiff",
        "docx", "txt", "pdf", "epub",
        "mp3", "wav", "ogg", "aac", "flac", "m4a",
        "mp4", "avi", "mov", "mkv", "webm"
    ]
)

if uploaded_file is not None:
    category, formats = get_output_formats(uploaded_file.name)

    if not formats:
        st.error("Unsupported file type.")
    else:
        st.success(f"Detected: {category.upper()} file")
        output_format = st.selectbox("Convert to", formats)

        if st.button("Convert"):
            with st.spinner("Converting..."):
                with tempfile.TemporaryDirectory() as tmpdir:
                    input_path = os.path.join(tmpdir, uploaded_file.name)

                    with open(input_path, "wb") as f:
                        f.write(uploaded_file.read())

                    try:
                        if category == "image":
                            output_path = convert_image(input_path, output_format)
                        elif category == "document":
                            output_path = convert_document(input_path, output_format)
                        elif category in ("audio", "video"):
                            output_path = convert_audio_video(input_path, output_format)
                        else:
                            raise ValueError("Unknown file category.")

                        with open(output_path, "rb") as f:
                            output_bytes = f.read()

                        out_name = os.path.basename(output_path)

                        st.success("Conversion complete.")
                        st.download_button(
                            label=f"Download {out_name}",
                            data=output_bytes,
                            file_name=out_name,
                            mime="application/octet-stream"
                        )

                    except Exception as e:
                        st.error(f"Conversion failed: {e}")